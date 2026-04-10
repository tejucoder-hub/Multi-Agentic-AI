
# # # # # # ----------------------------------------------------------------------------------------------------------


import io
import json
import logging
import os
import sys
import uuid
from datetime import datetime
from pathlib import Path
# from metrics_logger import load_metrics
import PyPDF2
import streamlit as st
from docx import Document
from langchain_core.messages import AIMessage
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import pandas as pd
CURRENT_DIR = Path(__file__).resolve().parent
if str(CURRENT_DIR) not in sys.path:
    sys.path.insert(0, str(CURRENT_DIR))

from ai_researcher_2 import (  # noqa: E402
    analyse_image_with_openai,
    build_graph,
    generate_image_with_openai,
    is_image_generation_request,
)
from adaptive_rag import run_adaptive_rag_workflow  # noqa: E402
from write_pdf import render_latex_pdf  # noqa: E402

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

CHAT_STORE_FILE = CURRENT_DIR / "chat_store.json"
EXPORT_DIR = CURRENT_DIR / "exports"
EXPORT_DIR.mkdir(exist_ok=True)


def default_store():
    return {
        "current_chat_id": None,
        "openai_model": "gpt-4.1",
        "openai_image_model": "gpt-image-1",
        "agent_mode": "standard",
        "chats": {},
    }


def load_chat_store():
    if not CHAT_STORE_FILE.exists():
        return default_store()

    try:
        with open(CHAT_STORE_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        defaults = default_store()
        for key, value in defaults.items():
            data.setdefault(key, value)
        return data
    except Exception as e:
        st.warning(f"Could not load previous chats: {e}")
        return default_store()


def save_chat_store():
    data = {
        "current_chat_id": st.session_state.current_chat_id,
        "openai_model": st.session_state.openai_model,
        "openai_image_model": st.session_state.openai_image_model,
        "agent_mode": st.session_state.agent_mode,
        "chats": st.session_state.chats,
    }
    with open(CHAT_STORE_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)


def make_chat_title(text: str) -> str:
    text = (text or "").strip()
    if not text:
        return "New Chat"
    return text[:48] + ("..." if len(text) > 48 else "")


def create_new_chat():
    chat_id = str(uuid.uuid4())
    st.session_state.chats[chat_id] = {
        "title": "New Chat",
        "messages": [],
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }
    st.session_state.current_chat_id = chat_id
    save_chat_store()


def get_current_chat():
    return st.session_state.chats[st.session_state.current_chat_id]


def delete_current_chat():
    current_id = st.session_state.current_chat_id
    if current_id in st.session_state.chats:
        del st.session_state.chats[current_id]

    if not st.session_state.chats:
        create_new_chat()
    else:
        st.session_state.current_chat_id = list(st.session_state.chats.keys())[0]
        save_chat_store()


def rename_chat(chat_id: str, new_title: str):
    if chat_id in st.session_state.chats and new_title.strip():
        st.session_state.chats[chat_id]["title"] = new_title.strip()
        save_chat_store()


def extract_pdf_text(uploaded_pdf) -> str:
    try:
        reader = PyPDF2.PdfReader(uploaded_pdf)
        text_parts = []
        for page in reader.pages[:20]:
            text_parts.append(page.extract_text() or "")
        return "\n".join(text_parts).strip()
    except Exception as e:
        return f"Could not read PDF: {e}"


def export_chat_as_txt(chat):
    lines = [f"Title: {chat['title']}", f"Created: {chat['created_at']}", "=" * 70]
    for msg in chat["messages"]:
        lines.append(f"\n{msg['role'].upper()}:\n{msg.get('content', '')}\n")
        if msg.get("image_path"):
            lines.append(f"Image: {msg['image_path']}\n")
        if msg.get("pdf_path"):
            lines.append(f"PDF: {msg['pdf_path']}\n")
    return "\n".join(lines)


def export_chat_as_json(chat):
    return json.dumps(chat, indent=2, ensure_ascii=False)


def export_chat_as_pdf_bytes(chat):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    _, height = A4
    y = height - 50

    def write_line(text):
        nonlocal y
        if y < 50:
            c.showPage()
            y = height - 50
        c.drawString(40, y, text[:110])
        y -= 15

    write_line(f"Title: {chat['title']}")
    write_line(f"Created: {chat['created_at']}")
    write_line("-" * 80)

    for msg in chat["messages"]:
        write_line(f"{msg['role'].upper()}:")
        for line in msg.get("content", "").splitlines():
            write_line(line)
        write_line("")

    c.save()
    buffer.seek(0)
    return buffer.getvalue()


def export_chat_as_docx_bytes(chat):
    doc = Document()
    doc.add_heading(chat["title"], level=1)
    doc.add_paragraph(f"Created: {chat['created_at']}")

    for msg in chat["messages"]:
        doc.add_heading(msg["role"].upper(), level=2)
        doc.add_paragraph(msg.get("content", ""))

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def extract_citations_from_text(text: str):
    lines = text.splitlines()
    citations = []
    for line in lines:
        if "http" in line or line.strip().startswith("["):
            citations.append(line.strip())
    return citations[:15]


def get_graph_bundle():
    graph_key = (
        st.session_state.openai_model,
        st.session_state.current_chat_id,
    )

    if st.session_state.get("_research_graph_cache_key") != graph_key:
        graph, config = build_graph(
            openai_model=st.session_state.openai_model,
            thread_id=st.session_state.current_chat_id,
        )
        st.session_state._research_graph_cache_key = graph_key
        st.session_state._research_graph_cache_value = (graph, config)

    return st.session_state._research_graph_cache_value


def get_file_size_label(file_path: str) -> str:
    try:
        size_bytes = os.path.getsize(file_path)
        if size_bytes < 1024:
            return f"{size_bytes} B"
        if size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.1f} KB"
        return f"{size_bytes / (1024 * 1024):.2f} MB"
    except Exception:
        return "Unknown size"


def render_pdf_attachment(file_path: str, key_suffix: str):
    if not file_path or not os.path.exists(file_path):
        return

    file_name = Path(file_path).name
    file_size = get_file_size_label(file_path)

    st.markdown(
        f"""
        <div style="
            margin-top: 0.6rem;
            padding: 0.9rem 1rem;
            border-radius: 16px;
            background: rgba(255,255,255,0.03);
            border: 1px solid rgba(255,255,255,0.08);
        ">
            <div style="
                font-size: 0.95rem;
                font-weight: 700;
                color: #f5f7fb;
                margin-bottom: 0.28rem;
                word-break: break-word;
            ">
                📄 {file_name}
            </div>
            <div style="
                font-size: 0.8rem;
                color: #aeb8c5;
            ">
                PDF document • {file_size}
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    with open(file_path, "rb") as f:
        pdf_bytes = f.read()

    st.download_button(
        label="Download file",
        data=pdf_bytes,
        file_name=file_name,
        mime="application/pdf",
        key=f"pdf_attachment_{key_suffix}",
        use_container_width=True,
    )


def render_message(msg, idx):
    with st.chat_message(msg["role"]):
        st.markdown(msg.get("content", ""))

        if msg.get("image_path") and os.path.exists(msg["image_path"]):
            st.image(msg["image_path"], use_container_width=True)

        pdf_path = msg.get("pdf_path")
        if pdf_path and os.path.exists(pdf_path):
            render_pdf_attachment(pdf_path, key_suffix=f"chat_{idx}")


def extract_ai_text(content) -> str:
    if isinstance(content, str):
        return content.strip()

    if isinstance(content, list):
        parts = []
        for item in content:
            if isinstance(item, dict) and item.get("type") == "text":
                text = (item.get("text") or "").strip()
                if text:
                    parts.append(text)
        return "\n".join(parts).strip()

    return ""


def sanitise_assistant_output(text: str) -> str:
    if not text:
        return text

    lowered = text.lower()

    bad_signals = [
        "you can download the pdf version using the link below",
        "download pdf:",
        "sandbox:/",
        "your research paper draft",
        "your research paper pdf has been generated successfully",
        "use the attached file below to download it",
        "if you need further customization",
        "if you need further customisation",
        "pdf version",
    ]

    if any(signal in lowered for signal in bad_signals):
        return "Your research paper is ready. The PDF file is attached below."

    return text


def normalise_export_content(text: str) -> str:
    if not text:
        return ""

    lines = text.strip().splitlines()
    filtered_lines = []

    bad_phrases = [
        "you can download the pdf version using the link below",
        "download pdf:",
        "sandbox:/",
        "your research paper draft",
        "your research paper pdf has been generated successfully",
        "use the attached file below to download it",
        "if you need further customization",
        "if you need further customisation",
        "pdf version",
    ]

    for line in lines:
        low = line.strip().lower()

        if not low:
            filtered_lines.append("")
            continue

        if any(phrase in low for phrase in bad_phrases):
            continue

        filtered_lines.append(line)

    return "\n".join(filtered_lines).strip()


def prepare_content_for_pdf(text: str) -> str:
    cleaned = normalise_export_content(text)
    return cleaned.strip() if cleaned else ""


def build_chat_history_messages(chat_messages: list[dict]) -> list[dict]:
    history = []
    for m in chat_messages:
        role = m.get("role")
        content = (m.get("content") or "").strip()
        if role in {"user", "assistant"} and content:
            history.append({"role": role, "content": content})
    return history


def looks_like_discovery_query(text: str) -> bool:
    lowered = (text or "").lower().strip()
    discovery_keywords = [
        "recent paper",
        "recent papers",
        "latest paper",
        "latest papers",
        "find papers",
        "show papers",
        "arxiv",
        "survey paper",
        "survey papers",
        "research papers",
        "recent research",
        "latest research",
        "paper on",
        "papers on",
        "find research paper",
        "give me papers",
    ]
    return any(k in lowered for k in discovery_keywords)


def looks_like_pdf_request(text: str) -> bool:
    lowered = (text or "").lower().strip()
    pdf_keywords = [
        "generate pdf",
        "make pdf",
        "create pdf",
        "export pdf",
        "export as pdf",
        "convert to pdf",
        "save as pdf",
        "pdf bana",
        "pdf banao",
        "make this pdf",
        "create pdf of this",
        "generate paper pdf",
        "in pdf format",
        "pdf format",
    ]
    return any(k in lowered for k in pdf_keywords)


def looks_like_draft_request(text: str) -> bool:
    lowered = (text or "").lower().strip()
    draft_keywords = [
        "write a research paper",
        "write research paper",
        "draft a research paper",
        "draft paper",
        "make a research paper",
        "generate a research paper",
        "write a paper on",
        "draft a paper on",
        "create a paper on",
        "prepare a research draft",
        "research draft",
    ]
    return any(k in lowered for k in draft_keywords)


def is_combined_draft_pdf_request(text: str) -> bool:
    return looks_like_draft_request(text) and looks_like_pdf_request(text)


def should_use_adaptive_rag(user_input: str) -> bool:
    if st.session_state.agent_mode != "adaptive_rag":
        return False
    return looks_like_discovery_query(user_input)


def stream_graph_answer(graph, config, chat_messages, response_box, status_box) -> str:
    final_text = ""

    for event in graph.stream({"messages": chat_messages}, config, stream_mode="values"):
        message = event["messages"][-1]

        if isinstance(message, AIMessage):
            text = extract_ai_text(message.content)
            if text:
                final_text = text
                response_box.markdown(sanitise_assistant_output(final_text))

            if getattr(message, "tool_calls", None):
                tools = [
                    tc.get("name", "tool")
                    for tc in message.tool_calls
                    if isinstance(tc, dict)
                ]
                if tools:
                    status_box.info(f"Using tools: {', '.join(tools)}")

    return final_text.strip()


def is_exportable_research_content(text: str) -> bool:
    if not text:
        return False

    cleaned = text.strip()
    if len(cleaned) < 80:
        return False

    lowered = cleaned.lower()

    strong_signals = [
        "abstract",
        "introduction",
        "methodology",
        "literature review",
        "discussion",
        "conclusion",
        "references",
        "research paper",
        "title",
        "background",
        "results",
    ]

    score = sum(1 for s in strong_signals if s in lowered)

    if score >= 1:
        return True

    if len(cleaned) >= 400:
        return True

    return False


def get_latest_exportable_assistant_content(chat) -> str:
    fallback = ""

    for msg in reversed(chat["messages"]):
        if msg.get("role") != "assistant":
            continue

        raw_export_content = prepare_content_for_pdf(
            (msg.get("raw_export_content") or "").strip()
        )
        if raw_export_content and is_exportable_research_content(raw_export_content):
            return raw_export_content

        content = prepare_content_for_pdf((msg.get("content") or "").strip())
        if not content:
            continue

        lowered = content.lower()

        if "error while running agent" in lowered:
            continue
        if "pdf has been generated successfully" in lowered:
            continue
        if "use the attached file below to download it" in lowered:
            continue

        if is_exportable_research_content(content):
            return content

        if len(content) > 300 and not fallback:
            fallback = content

    return fallback


def generate_pdf_from_latest_chat_draft(chat) -> str:
    latest_content = get_latest_exportable_assistant_content(chat)
    latest_content = prepare_content_for_pdf(latest_content)

    if not latest_content:
        raise ValueError(
            "No valid research draft found. Please first generate a proper research paper draft, then request PDF."
        )

    try:
        pdf_path = render_latex_pdf.invoke({"content": latest_content})
        return pdf_path
    except Exception as e:
        raise RuntimeError(f"PDF generation failed: {e}")


def format_adaptive_rag_response(result: dict) -> str:
    validated_query = (result.get("validated_query") or "").strip()
    refined_query = (result.get("refined_query") or "").strip()
    retrieval_queries = result.get("retrieval_queries", []) or []
    retrieved_docs = result.get("retrieved_documents", []) or []
    relevant_docs = result.get("relevant_documents", []) or []
    final_answer = (result.get("final_answer") or "").strip()
    confidence_note = (result.get("confidence_note") or "").strip()
    retry_count = result.get("retry_count", 0)
    retrieval_error = (result.get("retrieval_error") or "").strip()
    refinement_reason = (result.get("refinement_reason") or "").strip()

    def clean_text(text: str) -> str:
        return " ".join((text or "").split())

    def short_text(text: str, limit: int = 220) -> str:
        text = clean_text(text)
        if len(text) <= limit:
            return text
        return text[:limit].rstrip() + "..."

    def confidence_label(note: str) -> str:
        lowered = (note or "").lower()
        if "moderate confidence" in lowered:
            return "Moderate"
        if "high confidence" in lowered:
            return "High"
        if "preliminary" in lowered:
            return "Preliminary"
        if "low" in lowered:
            return "Low"
        return "Not stated"

    def retrieval_status(error_text: str, doc_count: int) -> str:
        if error_text and doc_count > 0:
            return "Partial success"
        if error_text and doc_count == 0:
            return "Failed"
        if doc_count > 0:
            return "Success"
        return "No documents"

    parts = []

    parts.append("## Adaptive RAG Research Report")
    parts.append("")
    parts.append("### Workflow Summary")
    parts.append(f"- **Validated Query:** {validated_query or 'Not available'}")
    parts.append(f"- **Refined Query:** {refined_query or 'No refinement applied'}")
    parts.append(f"- **Retry Count:** {retry_count}")
    parts.append(f"- **Retrieved Documents:** {len(retrieved_docs)}")
    parts.append(f"- **Relevant Documents:** {len(relevant_docs)}")
    parts.append(f"- **Retrieval Status:** {retrieval_status(retrieval_error, len(retrieved_docs))}")
    parts.append(f"- **Confidence:** {confidence_label(confidence_note)}")

    if refinement_reason:
        parts.append(f"- **Refinement Reason:** {refinement_reason}")

    if retrieval_error:
        parts.append(f"- **Retrieval Note:** {retrieval_error}")

    parts.append("")
    parts.append("### Retrieval Queries")

    if retrieval_queries:
        for idx, query in enumerate(retrieval_queries, start=1):
            parts.append(f"{idx}. {query}")
    else:
        parts.append("No retrieval query recorded.")

    parts.append("")
    parts.append("### Retrieved Docs")

    if retrieved_docs:
        for idx, doc in enumerate(retrieved_docs[:8], start=1):
            title = clean_text(doc.get("title", "Untitled"))
            authors = doc.get("authors", []) or []
            authors_text = ", ".join(authors[:4]) if authors else "Not available"
            if len(authors) > 4:
                authors_text += ", et al."
            published = (doc.get("published") or "")[:10] or "Unknown"
            source = doc.get("source", "Unknown")
            link = doc.get("link", "")
            abstract_preview = short_text(doc.get("abstract", ""), 240)

            parts.append(f"**{idx}. {title}**")
            parts.append(f"- Authors: {authors_text}")
            parts.append(f"- Published: {published}")
            parts.append(f"- Source: {source}")
            if link:
                parts.append(f"- Link: {link}")
            if abstract_preview:
                parts.append(f"- Preview: {abstract_preview}")
            parts.append("")
    else:
        parts.append("No documents retrieved.")
        parts.append("")

    parts.append("### Relevant Docs")

    if relevant_docs:
        for idx, doc in enumerate(relevant_docs[:6], start=1):
            title = clean_text(doc.get("title", "Untitled"))
            label = doc.get("relevance_label", "unknown")
            score = doc.get("relevance_score", 0)
            matched_terms = doc.get("matched_terms", []) or []
            abstract_preview = short_text(doc.get("abstract", ""), 220)

            why_line = "This paper appears aligned with the query based on title and abstract matching."
            if label == "highly relevant":
                why_line = "This paper is strongly aligned with the query and should be prioritised for deeper reading."
            elif label == "partially relevant":
                why_line = "This paper is reasonably aligned with the query and may provide useful supporting context."

            parts.append(f"**{idx}. {title}**")
            parts.append(f"- Relevance: {label}")
            parts.append(f"- Score: {score}")
            if matched_terms:
                parts.append(f"- Matched Terms: {', '.join(matched_terms[:8])}")
            parts.append(f"- Why it matters: {why_line}")
            if abstract_preview:
                parts.append(f"- Evidence Preview: {abstract_preview}")
            parts.append("")
    else:
        parts.append("No relevant documents found.")
        parts.append("")

    if relevant_docs:
        key_takeaways = []
        for doc in relevant_docs[:3]:
            abstract = clean_text(doc.get("abstract", ""))
            title = clean_text(doc.get("title", "Untitled"))
            first_sentence = abstract.split(". ")[0].strip()
            if first_sentence:
                key_takeaways.append(f"- {first_sentence.rstrip('.')}.")
            else:
                key_takeaways.append(f"- {title} is one of the more relevant retrieved papers.")

        parts.append("### Key Takeaways")
        parts.extend(key_takeaways)
        parts.append("")

    parts.append("### Final Answer")
    parts.append(final_answer or "No grounded answer returned.")
    parts.append("")
    parts.append("### Confidence")
    parts.append(confidence_note or "No confidence note available.")

    return "\n".join(parts).strip()


def run_research_app():
    st.title("Research Workspace")
    st.caption(
        "Welcome to the Research AI Agent. Explore papers, citations, PDFs, and research workflows in one focused workspace."
    )

    if "loaded_store" not in st.session_state:
        data = load_chat_store()
        st.session_state.chats = data["chats"]
        st.session_state.current_chat_id = data["current_chat_id"]
        st.session_state.openai_model = data["openai_model"]
        st.session_state.openai_image_model = data["openai_image_model"]
        st.session_state.agent_mode = data["agent_mode"]
        st.session_state.loaded_store = True

        if not st.session_state.chats:
            create_new_chat()
        elif (
            not st.session_state.current_chat_id
            or st.session_state.current_chat_id not in st.session_state.chats
        ):
            st.session_state.current_chat_id = list(st.session_state.chats.keys())[0]
            save_chat_store()

    st.markdown(
        """
        <style>
        .block-container {
            max-width: 1000px;
            padding-top: 1.2rem;
            padding-bottom: 1.2rem;
        }

        h1, h2, h3, h4, h5, h6, p, label, div, span {
            color: var(--text) !important;
        }

        section[data-testid="stSidebar"] {
            background: linear-gradient(180deg, #0b1017 0%, #090d14 100%);
            border-right: 1px solid rgba(255,255,255,0.08);
        }

        .status-box {
            padding: 0.9rem 1rem;
            border-radius: 16px;
            background: rgba(255,255,255,0.035);
            border: 1px solid rgba(255,255,255,0.08);
            margin-bottom: 1rem;
            box-shadow: 0 10px 20px rgba(0,0,0,0.14);
        }

        .research-hero {
            position: relative;
            padding: 1rem 1rem 1.05rem 1rem;
            margin-bottom: 1rem;
            border-radius: 20px;
            background:
                radial-gradient(circle at top right, rgba(76, 201, 240, 0.16), transparent 30%),
                linear-gradient(180deg, rgba(12,16,23,0.96) 0%, rgba(8,11,17,0.98) 100%);
            border: 1px solid rgba(120, 190, 255, 0.14);
            box-shadow: 0 10px 28px rgba(0,0,0,0.22);
            overflow: hidden;
        }

        .research-hero-badge {
            display: inline-block;
            padding: 0.22rem 0.55rem;
            border-radius: 999px;
            font-size: 0.72rem;
            font-weight: 700;
            letter-spacing: 0.04em;
            color: #9ed8ff !important;
            background: rgba(79, 172, 254, 0.10);
            border: 1px solid rgba(79, 172, 254, 0.18);
            margin-bottom: 0.7rem;
        }

        .research-hero-title {
            font-size: 1.35rem;
            font-weight: 800;
            line-height: 1.15;
            color: #f8fbff !important;
            margin-bottom: 0.45rem;
        }

        .research-hero-subtitle {
            font-size: 0.92rem;
            line-height: 1.55;
            color: #b8c2cf !important;
        }

        .research-overview-card {
            padding: 1rem;
            margin-bottom: 1rem;
            border-radius: 20px;
            background: rgba(255,255,255,0.025);
            border: 1px solid rgba(255,255,255,0.08);
            box-shadow: 0 8px 20px rgba(0,0,0,0.18);
        }

        .research-card-title {
            font-size: 0.95rem;
            font-weight: 700;
            color: #eef4ff !important;
            margin-bottom: 0.85rem;
        }

        .research-stat-grid {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 0.7rem;
        }

        .research-stat-box {
            padding: 0.72rem 0.75rem;
            border-radius: 16px;
            background: linear-gradient(180deg, rgba(255,255,255,0.04) 0%, rgba(255,255,255,0.02) 100%);
            border: 1px solid rgba(255,255,255,0.06);
        }

        .research-stat-label {
            font-size: 0.74rem;
            color: #8ea0b8 !important;
            margin-bottom: 0.22rem;
        }

        .research-stat-value {
            font-size: 0.96rem;
            font-weight: 700;
            color: #f4f8ff !important;
        }

        .upload-panel {
            padding: 1rem 1rem 0.35rem 1rem;
            border-radius: 18px;
            background: rgba(255,255,255,0.025);
            border: 1px solid rgba(255,255,255,0.08);
            margin-bottom: 1rem;
            box-shadow: 0 8px 20px rgba(0,0,0,0.16);
        }

        .upload-panel-title {
            font-size: 1rem;
            font-weight: 700;
            margin-bottom: 0.3rem;
            color: #f3f6fb !important;
        }

        .upload-panel-subtitle {
            color: #aab5c2 !important;
            font-size: 0.92rem;
            margin-bottom: 0.85rem;
            line-height: 1.55;
        }

        .input-status {
            padding: 0.8rem 0.9rem;
            border-radius: 14px;
            background: rgba(255,255,255,0.03);
            border: 1px solid rgba(255,255,255,0.07);
            margin-bottom: 1rem;
        }

        .input-status strong {
            color: #f5f8ff !important;
        }

        div[data-testid="stChatMessage"] {
            border-radius: 16px;
        }

        .stButton > button {
            border-radius: 12px !important;
        }

        .stDownloadButton > button {
            border-radius: 12px !important;
        }

        .stTextInput input,
        .stTextArea textarea {
            border-radius: 12px !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    with st.sidebar:
        st.markdown(
            """
            <div class="research-hero">
                <div class="research-hero-badge">Research Agent</div>
                <div class="research-hero-title">Research Workspace</div>
                <div class="research-hero-subtitle">
                    Focused environment for papers, citations, literature review, PDF understanding, and multimodal research workflows.
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        st.markdown(
            """
            <div class="research-overview-card">
                <div class="research-card-title">Workspace Overview</div>
                <div class="research-stat-grid">
                    <div class="research-stat-box">
                        <div class="research-stat-label">Mode</div>
                        <div class="research-stat-value">Research</div>
                    </div>
                    <div class="research-stat-box">
                        <div class="research-stat-label">Search</div>
                        <div class="research-stat-value">Active</div>
                    </div>
                    <div class="research-stat-box">
                        <div class="research-stat-label">Citations</div>
                        <div class="research-stat-value">Ready</div>
                    </div>
                    <div class="research-stat-box">
                        <div class="research-stat-label">PDF Analysis</div>
                        <div class="research-stat-value">Ready</div>
                    </div>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        st.markdown("## Agent Mode")
        agent_mode_label = st.select_slider(
            "Choose workflow",
            options=["Standard Research", "Adaptive RAG"],
            value="Adaptive RAG" if st.session_state.agent_mode == "adaptive_rag" else "Standard Research",
            key="research_agent_mode_slider",
            label_visibility="collapsed",
        )
        st.session_state.agent_mode = (
            "adaptive_rag" if agent_mode_label == "Adaptive RAG" else "standard"
        )
        save_chat_store()

        st.markdown("## Chats")

        cc1, cc2 = st.columns([3, 1])
        with cc1:
            if st.button("➕ New Chat", key="research_new_chat_btn", use_container_width=True):
                create_new_chat()
                st.rerun()
        with cc2:
            if st.button("🗑", key="research_delete_chat_btn", use_container_width=True):
                delete_current_chat()
                st.rerun()

        search_text = st.text_input("Search chats", key="research_search_chats")
        st.markdown("### Chat History")

        items = list(st.session_state.chats.items())[::-1]
        for chat_id, chat_data in items:
            title = chat_data.get("title", "Untitled")
            if search_text and search_text.lower() not in title.lower():
                continue

            if st.button(f"{title}", key=f"chat_{chat_id}", use_container_width=True):
                st.session_state.current_chat_id = chat_id
                save_chat_store()
                st.rerun()

            st.caption(chat_data.get("created_at", ""))

        st.markdown("---")

    current_chat = get_current_chat()

    with st.expander("Chat settings, exports, and citation panel"):
        rename_col1, rename_col2 = st.columns([3, 1])
        with rename_col1:
            new_title = st.text_input(
                "Rename current chat",
                value=current_chat["title"],
                key="research_rename_chat_input",
            )
        with rename_col2:
            st.write("")
            st.write("")
            if st.button("Save title", key="research_save_chat_title_btn", use_container_width=True):
                rename_chat(st.session_state.current_chat_id, new_title)
                st.rerun()

        export_tab, citation_tab = st.tabs(["Exports", "Citations"])

        with export_tab:
            txt_data = export_chat_as_txt(current_chat)
            json_data = export_chat_as_json(current_chat)
            pdf_data = export_chat_as_pdf_bytes(current_chat)
            docx_data = export_chat_as_docx_bytes(current_chat)

            e1, e2, e3, e4 = st.columns(4)
            with e1:
                st.download_button(
                    "TXT",
                    txt_data,
                    file_name="chat.txt",
                    mime="text/plain",
                    use_container_width=True,
                    key="research_download_txt",
                )
            with e2:
                st.download_button(
                    "JSON",
                    json_data,
                    file_name="chat.json",
                    mime="application/json",
                    use_container_width=True,
                    key="research_download_json",
                )
            with e3:
                st.download_button(
                    "PDF",
                    pdf_data,
                    file_name="chat.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    key="research_download_pdf",
                )
            with e4:
                st.download_button(
                    "DOCX",
                    docx_data,
                    file_name="chat.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="research_download_docx",
                )

        with citation_tab:
            all_citations = []
            for m in reversed(current_chat["messages"]):
                if m.get("role") == "assistant":
                    all_citations.extend(extract_citations_from_text(m.get("content", "")))
                    if len(all_citations) >= 15:
                        break

            if all_citations:
                for c in all_citations[:15]:
                    st.markdown(f"- {c}")
            else:
                st.caption("No citations detected yet.")

    status_info = (
        f"Mode: {'Adaptive RAG' if st.session_state.agent_mode == 'adaptive_rag' else 'Standard Research'} | "
        f"Current chat: {current_chat['title']}"
    )
    st.markdown(f'<div class="status-box">{status_info}</div>', unsafe_allow_html=True)

    try:
        graph, config = get_graph_bundle()
    except Exception as e:
        st.error(f"Model setup failed: {e}")
        st.stop()

    st.markdown(
        """
        <div class="upload-panel">
            <div class="upload-panel-title">Research inputs</div>
            <div class="upload-panel-subtitle">
                Upload a PDF for summarisation or an image for multimodal analysis. If neither is attached, the agent will continue with standard research chat.
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    up1, up2 = st.columns(2)

    with up1:
        uploaded_pdf = st.file_uploader(
            "Upload PDF for summary",
            type=["pdf"],
            accept_multiple_files=False,
            key="research_pdf_uploader",
        )

    with up2:
        uploaded_image = st.file_uploader(
            "Upload image",
            type=["png", "jpg", "jpeg", "webp"],
            accept_multiple_files=False,
            key="research_image_uploader",
        )

    active_input_parts = []
    if uploaded_pdf is not None:
        active_input_parts.append(f"PDF attached: {uploaded_pdf.name}")
    if uploaded_image is not None:
        active_input_parts.append(f"Image attached: {uploaded_image.name}")

    if active_input_parts:
        st.markdown(
            f"""
            <div class="input-status">
                <strong>Active input:</strong> {" | ".join(active_input_parts)}
            </div>
            """,
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            """
            <div class="input-status">
                <strong>Active input:</strong> None. The agent will respond as a standard research assistant.
            </div>
            """,
            unsafe_allow_html=True,
        )

    if uploaded_image is not None:
        st.image(uploaded_image, caption="Uploaded image", use_container_width=True)

    for idx, msg in enumerate(current_chat["messages"]):
        render_message(msg, idx)


    



    user_input = st.chat_input(
        "Ask for papers, summaries, web research, PDF analysis, or image generation..."
    )

    if user_input:
        current_chat = get_current_chat()

        if current_chat["title"] == "New Chat":
            current_chat["title"] = make_chat_title(user_input)

        user_message = {"role": "user", "content": user_input}
        current_chat["messages"].append(user_message)
        st.session_state.chats[st.session_state.current_chat_id] = current_chat
        save_chat_store()

        with st.chat_message("user"):
            st.markdown(user_input)

        assistant_message = {"role": "assistant", "content": ""}

        with st.chat_message("assistant"):
            response_box = st.empty()
            status_box = st.empty()
            attachment_container = st.container()

            try:
                if is_combined_draft_pdf_request(user_input):
                    status_box.info("Generating research draft...")

                    chat_messages = build_chat_history_messages(current_chat["messages"])

                    final_text = stream_graph_answer(
                        graph=graph,
                        config=config,
                        chat_messages=chat_messages,
                        response_box=response_box,
                        status_box=status_box,
                    )

                    raw_draft = prepare_content_for_pdf((final_text or "").strip())

                    if not raw_draft or not is_exportable_research_content(raw_draft):
                        raise ValueError("The generated response was not a valid research paper draft.")

                    assistant_message["content"] = "Your research paper is ready. The PDF file is attached below."
                    assistant_message["raw_export_content"] = raw_draft
                    response_box.markdown(assistant_message["content"])

                    status_box.info("Generating PDF from the new draft...")
                    pdf_path = render_latex_pdf.invoke({"content": raw_draft})
                    assistant_message["pdf_path"] = pdf_path

                    with attachment_container:
                        if os.path.exists(pdf_path):
                            render_pdf_attachment(pdf_path, key_suffix=f"current_{uuid.uuid4().hex}")
                        else:
                            st.error("PDF file was generated but could not be found for download.")

                    status_box.empty()

                elif looks_like_pdf_request(user_input):
                    status_box.info("Generating PDF from latest draft...")
                    pdf_path = generate_pdf_from_latest_chat_draft(current_chat)

                    assistant_message["content"] = "Your research paper PDF is ready. The file is attached below."
                    assistant_message["pdf_path"] = pdf_path
                    response_box.markdown(assistant_message["content"])

                    with attachment_container:
                        if os.path.exists(pdf_path):
                            render_pdf_attachment(pdf_path, key_suffix=f"current_{uuid.uuid4().hex}")
                        else:
                            st.error("PDF file was generated but could not be found for download.")

                    status_box.empty()

                elif uploaded_pdf is not None:
                    status_box.info("Reading PDF...")
                    pdf_text = extract_pdf_text(uploaded_pdf)

                    pdf_prompt = f"""The user uploaded a PDF.

User request:
{user_input}

PDF extracted text:
{pdf_text[:15000]}

Please summarise it clearly, identify key contributions, methods, findings, limitations, and useful next steps."""

                    final_text = stream_graph_answer(
                        graph=graph,
                        config=config,
                        chat_messages=[{"role": "user", "content": pdf_prompt}],
                        response_box=response_box,
                        status_box=status_box,
                    )

                    raw_text = final_text or "No PDF summary returned."
                    assistant_message["content"] = sanitise_assistant_output(raw_text)
                    assistant_message["raw_export_content"] = prepare_content_for_pdf(raw_text)
                    status_box.empty()

                elif uploaded_image is not None:
                    status_box.info("Analysing image...")
                    answer = analyse_image_with_openai(
                        prompt=user_input,
                        image_bytes=uploaded_image.getvalue(),
                        mime_type=uploaded_image.type or "image/png",
                        model=st.session_state.openai_model,
                    )
                    raw_text = answer or ""
                    display_text = sanitise_assistant_output(raw_text)
                    assistant_message["content"] = display_text
                    assistant_message["raw_export_content"] = prepare_content_for_pdf(raw_text)
                    response_box.markdown(display_text)
                    status_box.empty()

                elif is_image_generation_request(user_input):
                    status_box.info("Generating image...")
                    image_path = generate_image_with_openai(
                        prompt=user_input,
                        model=st.session_state.openai_image_model,
                    )
                    assistant_message["content"] = "I generated an image for your request."
                    assistant_message["image_path"] = image_path
                    response_box.markdown(assistant_message["content"])
                    st.image(image_path, use_container_width=True)
                    status_box.empty()

                else:
                    if should_use_adaptive_rag(user_input):
                        status_box.info("Running Adaptive RAG workflow...")
                        result = run_adaptive_rag_workflow(user_input)
                        raw_text = format_adaptive_rag_response(result)
                        display_text = sanitise_assistant_output(raw_text)
                        assistant_message["content"] = display_text
                        assistant_message["raw_export_content"] = prepare_content_for_pdf(raw_text)
                        response_box.markdown(display_text)
                        status_box.empty()
                    else:
                        status_box.info("Thinking...")
                        chat_messages = build_chat_history_messages(current_chat["messages"])

                        final_text = stream_graph_answer(
                            graph=graph,
                            config=config,
                            chat_messages=chat_messages,
                            response_box=response_box,
                            status_box=status_box,
                        )

                        raw_text = final_text or "No response returned."
                        assistant_message["content"] = sanitise_assistant_output(raw_text)
                        assistant_message["raw_export_content"] = prepare_content_for_pdf(raw_text)
                        status_box.empty()

            except Exception as e:
                assistant_message["content"] = f"Error while running agent: {e}"
                response_box.error(assistant_message["content"])
                logger.exception("Agent execution failed")
                status_box.empty()

        current_chat["messages"].append(assistant_message)
        st.session_state.chats[st.session_state.current_chat_id] = current_chat
        save_chat_store()